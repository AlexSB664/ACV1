# pip3 install python-docx
from docx import Document
import re
file_path = '/home/alex/WEB/Contrato.docx'
regex1=re.compile(r'2019')
replace1 =r'2018'
def docx_replace_regex(doc_obj, regex , replace):
	for p in doc_obj.paragraphs:
		if regex.search(p.text):
			inline = p.runs
			for i in range(len(inline)):
				if regex.search(inline[i].text):
					print(inline[i].text)
					inline[i].underline = True
					inline[i].text = replace
	for table in doc_obj.tables:
		for row in table.rows:
			for cell in row.cells:
				docx_replace_regex(cell, regex , replace)
document = Document(file_path)
docx_replace_regex(document,regex1,replace1)
document.save('xdxdxd.docx')
